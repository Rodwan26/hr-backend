import os
import PyPDF2
import docx
from typing import List, Dict, Optional
from sqlalchemy.orm import Session
from fastapi import UploadFile
from app.models.document import Document
from app.models.document_chunk import DocumentChunk
from app.models.company import Company
from app.services.embedding_service import generate_embeddings, hybrid_search
import hashlib
import re
import logging
import time

logger = logging.getLogger(__name__)

# Configuration
CHUNK_SIZE = 800
CHUNK_OVERLAP = 200
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.csv'}

def validate_file(file: UploadFile) -> tuple[bool, str]:
    """Validate uploaded file."""
    # Check extension
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        return False, f"File type {file_ext} not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
    
    # Check file size (will be checked after reading)
    return True, ""

def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from uploaded file based on type."""
    text = ""
    
    try:
        if file_type == '.pdf':
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        
        elif file_type == '.docx':
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        
        elif file_type in ['.txt', '.csv']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from file: {str(e)}")

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks - FAST version."""
    logger.info(f"Starting chunking - text length: {len(text)} chars")
    
    if len(text) <= chunk_size:
        logger.info("Text is small, returning as single chunk")
        return [text]
    
    chunks = []
    start = 0
    chunk_count = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Simple break - no complex sentence boundary detection
        if end < len(text):
            # Just find nearest space
            space_pos = text.find(' ', end)
            if space_pos != -1 and space_pos < end + 100:
                end = space_pos
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
            chunk_count += 1
            if chunk_count % 10 == 0:
                logger.info(f"Created {chunk_count} chunks...")
        
        # Move start forward with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    logger.info(f"Chunking complete - created {len(chunks)} chunks")
    return chunks

async def process_uploaded_file(
    file: UploadFile,
    company_id: int,
    uploaded_by: str,
    db: Session,
    upload_dir: str = "uploads"
) -> Document:
    """
    Process uploaded file: extract text, chunk, generate embeddings, and store.
    """
    process_start = time.time()
    logger.info(f"=== Starting file upload process ===")
    logger.info(f"Filename: {file.filename}")
    logger.info(f"Company ID: {company_id}, Uploaded by: {uploaded_by}")
    
    try:
        # Step 1: Validate file
        logger.info("Step 1: Validating file...")
        is_valid, error_msg = validate_file(file)
        if not is_valid:
            logger.error(f"File validation failed: {error_msg}")
            raise ValueError(error_msg)
        logger.info("Step 1: File validation passed")
        
        # Step 2: Create upload directory
        logger.info("Step 2: Creating upload directory...")
        company_upload_dir = os.path.join(upload_dir, str(company_id))
        os.makedirs(company_upload_dir, exist_ok=True)
        logger.info(f"Step 2: Upload directory ready: {company_upload_dir}")
        
        # Step 3: Prepare file path
        file_ext = os.path.splitext(file.filename)[1].lower()
        file_hash = hashlib.md5(f"{file.filename}{company_id}".encode()).hexdigest()
        file_path = os.path.join(company_upload_dir, f"{file_hash}{file_ext}")
        logger.info(f"Step 3: File path prepared: {file_path}")
        
        # Step 4: Read and save file
        logger.info("Step 4: Reading file content...")
        read_start = time.time()
        content = await file.read()
        file_size = len(content)
        read_time = time.time() - read_start
        logger.info(f"Step 4: File read complete - Size: {file_size} bytes ({file_size/(1024*1024):.2f}MB) in {read_time:.2f}s")
        
        if file_size > MAX_FILE_SIZE:
            logger.error(f"File size {file_size} exceeds limit {MAX_FILE_SIZE}")
            raise ValueError(f"File size exceeds {MAX_FILE_SIZE / (1024*1024)}MB limit")
        
        logger.info("Step 5: Saving file to disk...")
        save_start = time.time()
        with open(file_path, 'wb') as f:
            f.write(content)
        save_time = time.time() - save_start
        logger.info(f"Step 5: File saved in {save_time:.2f}s")
    
        # Step 6: Extract text
        logger.info("Step 6: Extracting text from file...")
        extract_start = time.time()
        try:
            text = extract_text_from_file(file_path, file_ext)
            extract_time = time.time() - extract_start
            logger.info(f"Step 6: Text extraction complete - Length: {len(text)} chars in {extract_time:.2f}s")
        except Exception as e:
            logger.error(f"Step 6: Text extraction failed: {e}", exc_info=True)
            if os.path.exists(file_path):
                os.remove(file_path)
            raise ValueError(f"Error extracting text from file: {str(e)}")
        
        if not text or not text.strip():
            logger.error("Step 6: Extracted text is empty")
            if os.path.exists(file_path):
                os.remove(file_path)
            raise ValueError("Could not extract text from file or file is empty")
        
        if len(text.strip()) < 10:
            logger.error(f"Step 6: Extracted text too short: {len(text.strip())} chars")
            if os.path.exists(file_path):
                os.remove(file_path)
            raise ValueError("Extracted text is too short (minimum 10 characters required)")
        
        # Step 7: Create document record
        logger.info("Step 7: Creating document record in database...")
        doc_start = time.time()
        document = Document(
            filename=file.filename,
            file_path=file_path,
            file_type=file_ext,
            uploaded_by=uploaded_by,
            company_id=company_id
        )
        db.add(document)
        db.flush()
        doc_time = time.time() - doc_start
        logger.info(f"Step 7: Document record created (ID: {document.id}) in {doc_time:.2f}s")
    
        # Step 8: Chunk text
        logger.info("Step 8: Chunking text...")
        chunk_start = time.time()
        chunks = chunk_text(text)
        chunk_time = time.time() - chunk_start
        logger.info(f"Step 8: Created {len(chunks)} chunks in {chunk_time:.2f}s")
        
        if not chunks or len(chunks) == 0:
            logger.error("Step 8: No chunks created")
            raise ValueError("Failed to create text chunks from document")
        
        valid_chunks = [chunk for chunk in chunks if chunk and chunk.strip()]
        if not valid_chunks:
            logger.error("Step 8: All chunks are empty")
            raise ValueError("All text chunks are empty")
        
        logger.info(f"Step 8: Valid chunks: {len(valid_chunks)}")
        
        # Step 9: Generate embeddings
        logger.info("Step 9: Generating embeddings for chunks...")
        embed_start = time.time()
        try:
            embeddings = generate_embeddings(valid_chunks, db)
            embed_time = time.time() - embed_start
            logger.info(f"Step 9: Generated {len(embeddings)} embeddings in {embed_time:.2f}s")
        except Exception as embed_error:
            logger.error(f"Step 9: Embedding generation error: {embed_error}", exc_info=True)
            from app.services.embedding_service import _fallback_embedding
            embeddings = []
            for chunk in valid_chunks:
                try:
                    embeddings.append(_fallback_embedding(chunk))
                except Exception as e:
                    logger.warning(f"Fallback embedding failed: {e}, using zero vector")
                    embeddings.append([0.0] * 384)
        
        if not embeddings or len(embeddings) == 0:
            logger.error("Step 9: No embeddings generated")
            raise ValueError("Failed to generate embeddings for document chunks")
        
        if len(embeddings) != len(valid_chunks):
            logger.warning(f"Step 9: Embedding count mismatch: {len(embeddings)} vs {len(valid_chunks)}, fixing...")
            from app.services.embedding_service import _fallback_embedding
            if len(embeddings) < len(valid_chunks):
                for i in range(len(embeddings), len(valid_chunks)):
                    embeddings.append(_fallback_embedding(valid_chunks[i]))
            else:
                embeddings = embeddings[:len(valid_chunks)]
            logger.info(f"Step 9: Fixed embedding count to {len(embeddings)}")
        
        # Step 10: Create chunk records
        logger.info("Step 10: Creating chunk records in database...")
        chunk_record_start = time.time()
        chunk_count = 0
        for idx, (chunk_content, embedding) in enumerate(zip(valid_chunks, embeddings)):
            if not chunk_content or not chunk_content.strip():
                continue
            
            if not embedding or len(embedding) == 0:
                from app.services.embedding_service import _fallback_embedding
                try:
                    embedding = _fallback_embedding(chunk_content)
                except Exception as e:
                    logger.warning(f"Fallback embedding for chunk {idx} failed: {e}")
                    embedding = [0.0] * 384
            
            chunk = DocumentChunk(
                document_id=document.id,
                chunk_text=chunk_content,
                chunk_index=idx,
                embedding_vector=embedding
            )
            db.add(chunk)
            chunk_count += 1
            
            if chunk_count % 20 == 0:
                logger.info(f"Step 10: Created {chunk_count}/{len(valid_chunks)} chunk records...")
        
        chunk_record_time = time.time() - chunk_record_start
        logger.info(f"Step 10: Created {chunk_count} chunk records in {chunk_record_time:.2f}s")
        
        # Step 11: Commit to database
        if chunk_count > 0:
            logger.info("Step 11: Committing transaction to database...")
            commit_start = time.time()
            try:
                db.commit()
                commit_time = time.time() - commit_start
                logger.info(f"Step 11: Database commit successful in {commit_time:.2f}s")
                
                db.refresh(document)
                total_time = time.time() - process_start
                logger.info(f"=== Upload process completed successfully ===")
                logger.info(f"Document ID: {document.id}")
                logger.info(f"Total chunks: {chunk_count}")
                logger.info(f"Total processing time: {total_time:.2f}s")
                return document
            except Exception as commit_error:
                logger.error(f"Step 11: Database commit failed: {commit_error}", exc_info=True)
                db.rollback()
                raise ValueError(f"Database commit failed: {str(commit_error)}")
        else:
            logger.error("Step 11: No chunks to commit")
            raise ValueError("No valid chunks were created")
    
    except Exception as e:
        total_time = time.time() - process_start
        logger.error(f"=== Upload process FAILED after {total_time:.2f}s ===")
        logger.error(f"Error: {str(e)}", exc_info=True)
        
        try:
            db.rollback()
            logger.info("Database transaction rolled back")
        except Exception as rollback_error:
            logger.error(f"Rollback failed: {rollback_error}")
        
        if 'document' in locals() and document.id:
            try:
                logger.info(f"Cleaning up document record {document.id}")
                db.delete(document)
                db.commit()
            except Exception as cleanup_error:
                logger.error(f"Document cleanup failed: {cleanup_error}")
                db.rollback()
        
        if 'file_path' in locals() and os.path.exists(file_path):
            try:
                logger.info(f"Removing file: {file_path}")
                os.remove(file_path)
            except Exception as file_cleanup_error:
                logger.error(f"File cleanup failed: {file_cleanup_error}")
        
        error_msg = str(e)
        logger.error(f"Raising error: {error_msg}")
        raise ValueError(f"Error processing document: {error_msg}")

def query_documents(
    question: str,
    company_id: int,
    db: Session,
    top_k: int = 5,
    document_ids: Optional[List[int]] = None
) -> Dict:
    """Query documents using RAG."""
    from app.services.embedding_service import generate_embeddings, hybrid_search
    
    query_embeddings = generate_embeddings([question], db)
    query_embedding = query_embeddings[0] if query_embeddings else None
    
    if not query_embedding:
        return {
            "answer": "Unable to process query.",
            "sources": [],
            "confidence": 0.0
        }
    
    results = hybrid_search(question, query_embedding, company_id, db, top_k)
    
    if not results:
        return {
            "answer": "No relevant information found in documents.",
            "sources": [],
            "confidence": 0.0
        }
    
    from app.models.document import Document
    sources = []
    context_chunks = []
    
    for result in results:
        doc = db.query(Document).filter(Document.id == result["document_id"]).first()
        if doc:
            sources.append({
                "document_id": doc.id,
                "filename": doc.filename,
                "chunk_index": result["chunk_index"],
                "similarity": result.get("similarity", 0.0),
                "combined_score": result.get("combined_score", 0.0)
            })
            context_chunks.append(result["chunk_text"])
    
    answer, confidence = generate_rag_answer(question, context_chunks)
    
    return {
        "answer": answer,
        "sources": sources,
        "confidence": confidence
    }

def generate_rag_answer(question: str, context_chunks: List[str]) -> tuple[str, float]:
    """Generate answer using RAG with OpenRouter."""
    from app.services.openrouter_client import call_openrouter
    
    context = "\n\n".join([
        f"[Document Chunk {i+1}]:\n{chunk}"
        for i, chunk in enumerate(context_chunks)
    ])
    
    messages = [
        {
            "role": "system",
            "content": "You are a helpful assistant that answers questions based on provided document context. Always cite which document chunk you used. If the context doesn't contain the answer, say so clearly."
        },
        {
            "role": "user",
            "content": f"Context from documents:\n\n{context}\n\nQuestion: {question}\n\nProvide a clear answer based on the context above. If the answer is not in the context, state that clearly."
        }
    ]
    
    try:
        answer = call_openrouter(messages, temperature=0.3)
        confidence = min(1.0, len(answer) / 200.0)
        return answer, confidence
    except Exception as e:
        return f"Error generating answer: {str(e)}", 0.0

def delete_document(document_id: int, company_id: int, db: Session) -> bool:
    """Delete a document and all its chunks."""
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.company_id == company_id
    ).first()
    
    if not document:
        return False
    
    if os.path.exists(document.file_path):
        try:
            os.remove(document.file_path)
        except:
            pass
    
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
    db.delete(document)
    db.commit()
    
    return True